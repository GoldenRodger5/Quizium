import json
import sys
from pathlib import Path
import PyPDF2
import anthropic
import random
import os
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.responses import HTMLResponse
from pydantic import BaseModel
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse, parse_qs
import re
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not installed. DOC/DOCX support disabled. Install with: pip install python-docx")

# Load environment variables from .env file
load_dotenv()

# Claude API configuration
CLAUDE_API_KEY = os.getenv("CLAUDE_API_KEY")
if not CLAUDE_API_KEY:
    raise ValueError("CLAUDE_API_KEY not found in environment variables. Please check your .env file.")

client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)

app = FastAPI()

# Configure templates
templates = Jinja2Templates(directory="templates")

# Pydantic models for API requests
class URLRequest(BaseModel):
    url: str

# Add this to handle HEAD requests
@app.head("/")
async def head_root():
    return {}

@app.get("/")
async def root():
    return {"message": "Flashcard API is running"}

@app.get("/app", response_class=HTMLResponse)
async def get_app():
    """Serve the flashcard web application."""
    try:
        with open("templates/index.html", "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Application template not found")

def extract_text_from_pdf(pdf_path):
    """Extract text content from a PDF file."""
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None

def extract_text_from_url(url):
    """Extract text content from various web sources."""
    try:
        # Determine content type based on URL
        if 'youtube.com' in url or 'youtu.be' in url:
            return extract_youtube_transcript(url)
        elif 'wikipedia.org' in url:
            return extract_wikipedia_content(url)
        else:
            return extract_general_webpage(url)
    except Exception as e:
        error_msg = f"Error extracting content from URL: {e}"
        print(error_msg)
        raise Exception(error_msg)

def extract_general_webpage(url):
    """Extract text from general web pages."""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }
    
    response = requests.get(url, headers=headers, timeout=10)
    response.raise_for_status()
    
    soup = BeautifulSoup(response.content, 'html.parser')
    
    # Remove script and style elements
    for script in soup(["script", "style", "nav", "header", "footer", "aside"]):
        script.decompose()
    
    # Get text from main content areas
    main_content = soup.find('main') or soup.find('article') or soup.find('div', class_=re.compile(r'content|article|post'))
    
    if main_content:
        text = main_content.get_text()
    else:
        text = soup.get_text()
    
    # Clean up text
    lines = (line.strip() for line in text.splitlines())
    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
    text = ' '.join(chunk for chunk in chunks if chunk)
    
    return text

def extract_wikipedia_content(url):
    """Extract content from Wikipedia articles with better structure."""
    # Convert to API format for cleaner extraction
    page_title = url.split('/')[-1]
    api_url = f"https://en.wikipedia.org/api/rest_v1/page/summary/{page_title}"
    
    try:
        response = requests.get(api_url, timeout=10)
        if response.status_code == 200:
            data = response.json()
            return data.get('extract', '') + '\n\n' + extract_general_webpage(url)
    except:
        pass
    
    return extract_general_webpage(url)

def extract_youtube_transcript(url):
    """Extract transcript from YouTube videos using youtube-transcript-api."""
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        
        # Extract video ID
        if 'youtu.be' in url:
            video_id = url.split('/')[-1].split('?')[0]
        else:
            parsed_url = urlparse(url)
            if 'v' not in parse_qs(parsed_url.query):
                raise ValueError("Invalid YouTube URL: missing video ID")
            video_id = parse_qs(parsed_url.query)['v'][0]
        
        print(f"Attempting to extract transcript for video ID: {video_id}")
        
        # Get transcript
        transcript = YouTubeTranscriptApi.get_transcript(video_id)
        
        # Combine transcript text
        text = ' '.join([entry['text'] for entry in transcript])
        print(f"Successfully extracted {len(text)} characters from YouTube transcript")
        return text
        
    except ImportError:
        error_msg = "youtube-transcript-api not installed. Install with: pip install youtube-transcript-api"
        print(error_msg)
        raise Exception(error_msg)
    except Exception as e:
        error_msg = f"Error extracting YouTube transcript: {e}"
        print(error_msg)
        raise Exception(error_msg)

def generate_flashcards(text_content):
    """Generate flashcards from text using Claude AI."""
    
    # Limit input text length to prevent token overflow
    MAX_INPUT_LENGTH = 100000
    if len(text_content) > MAX_INPUT_LENGTH:
        print(f"⚠️  Input text is {len(text_content)} characters. Truncating to {MAX_INPUT_LENGTH} characters for better processing.")
        text_content = text_content[:MAX_INPUT_LENGTH] + "..."
    
    prompt = f"""
    Please analyze the following text and create comprehensive flashcards for studying. 
    Generate the following types of flashcards:
    1. Question-Answer pairs for key concepts
    2. Vocabulary terms with definitions
    3. Important facts and figures
    4. Any other relevant study material

    'category' is used to organize flashcards based on the category of the content, such as 'ai principles', or 'super cars'. The category should be specific to each flashcard's content.

    'difficulty' is used to indicate the complexity of the flashcard, such as 'easy', 'medium', or 'hard'.
    
    Please ensure the flashcards are well-structured and informative. Do not include any unnecessary information in the flashcards.

    When you generate flashcards, ensure all the information within the text is accounted for and included in the flashcards.

    Additionally, full sentences aren't required for the flashcards. The flashcards should be concise and to the point and clearly convey the information needed for effective studying.

    When generating the question, please make it clear the question being asked, such as "What is...", "Define...", or "Explain...".

    IMPORTANT: Return ONLY valid JSON with NO additional text, comments, or explanations. Format your response as:

    {{
        "flashcards": [
            {{
                "category": "category_name",
                "difficulty": "easy",
                "type": "question_answer",
                "question": "What is...",
                "answer": "..."
            }},
            {{
                "category": "category_name",
                "difficulty": "medium",
                "type": "vocabulary",
                "term": "...",
                "definition": "..."
            }},
            {{
                "category": "category_name",
                "difficulty": "hard",
                "type": "fact",
                "prompt": "...",
                "content": "..."
            }}
        ]
    }}

    Text to analyze:
    {text_content}
    """
    
    try:
        response = client.messages.create(
            model="claude-3-7-sonnet-20250219",
            max_tokens=8000,  # Increased from 4000 to 8000
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        
        # Extract JSON from response
        response_text = response.content[0].text
        print(f"Raw Claude response length: {len(response_text)} characters")
        print(f"Raw Claude response preview: {response_text[:200]}...")  # Reduced debug output
        
        # Try to find and parse JSON in the response
        start_idx = response_text.find('{')
        end_idx = response_text.rfind('}') + 1
        
        if start_idx != -1 and end_idx != 0:
            json_str = response_text[start_idx:end_idx]
            
            # Clean up common JSON formatting issues
            json_str = json_str.replace('\n', ' ').replace('\r', ' ')
            
            # Try to parse JSON with better error handling
            try:
                parsed_json = json.loads(json_str)
                print(f"✅ Successfully parsed JSON with {len(parsed_json.get('flashcards', []))} flashcards")
                return parsed_json
            except json.JSONDecodeError as e:
                print(f"JSON parsing error: {e}")
                print(f"Error at position {e.pos}")
                
                # Try to fix common issues and parse again
                try:
                    # Remove trailing commas before closing brackets/braces
                    import re
                    json_str = re.sub(r',(\s*[}\]])', r'\1', json_str)
                    
                    # Try to find incomplete JSON and fix it
                    if json_str.count('{') > json_str.count('}'):
                        # Add missing closing braces
                        missing_braces = json_str.count('{') - json_str.count('}')
                        json_str += '}' * missing_braces
                        print(f"🔧 Added {missing_braces} missing closing braces")
                    
                    if json_str.count('[') > json_str.count(']'):
                        # Add missing closing brackets
                        missing_brackets = json_str.count('[') - json_str.count(']')
                        json_str += ']' * missing_brackets
                        print(f"🔧 Added {missing_brackets} missing closing brackets")
                    
                    parsed_json = json.loads(json_str)
                    print(f"✅ Successfully fixed and parsed JSON with {len(parsed_json.get('flashcards', []))} flashcards")
                    return parsed_json
                    
                except json.JSONDecodeError as e2:
                    print(f"Failed to fix JSON: {e2}")
                    
                    # Last resort: try to extract partial flashcards
                    try:
                        # Look for individual flashcard objects
                        import re
                        flashcard_pattern = r'\{[^{}]*"type"[^{}]*\}'
                        matches = re.findall(flashcard_pattern, json_str)
                        
                        if matches:
                            print(f"🔧 Attempting to extract {len(matches)} partial flashcards")
                            flashcards = []
                            for match in matches[:10]:  # Limit to 10 flashcards
                                try:
                                    flashcard = json.loads(match)
                                    flashcards.append(flashcard)
                                except:
                                    continue
                            
                            if flashcards:
                                print(f"✅ Extracted {len(flashcards)} partial flashcards")
                                return {"flashcards": flashcards}
                    
                    except Exception as e3:
                        print(f"Partial extraction failed: {e3}")
                    
                    print("❌ All JSON parsing attempts failed")
                    return None
        else:
            print("Could not find valid JSON in response")
            return None
            
    except Exception as e:
        print(f"Error generating flashcards: {e}")
        return None

def save_flashcards(flashcards, output_path):
    """Save flashcards to a JSON file."""
    try:
        with open(output_path, 'w', encoding='utf-8') as file:
            json.dump(flashcards, file, indent=2, ensure_ascii=False)
        print(f"Flashcards saved to: {output_path}")
    except Exception as e:
        print(f"Error saving flashcards: {e}")

def load_flashcards(json_path):
    """Load flashcards from a JSON file."""
    try:
        with open(json_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
            return data.get('flashcards', [])
    except Exception as e:
        print(f"Error loading flashcards: {e}")
        return None

def check_answer(question, correct_answer, user_answer):
    """Check if user's answer is correct using Claude with more flexible evaluation."""
    prompt = f"""
    You are evaluating a student's answer to a study question. Be fair and flexible in your assessment.
    
    Question: {question}
    Expected Answer: {correct_answer}
    Student's Answer: {user_answer}
    
    Consider the student's answer CORRECT if:
    - It demonstrates understanding of the core concept
    - Contains the essential information, even if worded differently
    - Is factually accurate in relation to the question
    - Shows the student grasped the main point
    
    Consider INCORRECT only if:
    - The answer is factually wrong
    - Shows fundamental misunderstanding
    - Completely misses the point of the question
    - Contains significant errors
    
    Use your knowledge to evaluate if the student's answer is reasonable and correct, even if it doesn't match the expected answer word-for-word. Different correct explanations or phrasings should be accepted.
    
    Respond with ONLY: "CORRECT" or "INCORRECT"
    """
    
    try:
        response = client.messages.create(
            model="claude-3-7-sonnet-20250219",
            max_tokens=20,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        response_text = response.content[0].text.strip().upper()
        return "CORRECT" in response_text and "INCORRECT" not in response_text
    except:
        # Improved fallback logic
        user_lower = user_answer.lower().strip()
        correct_lower = correct_answer.lower().strip()
        
        # Exact match
        if user_lower == correct_lower:
            return True
            
        # Check if user answer contains key concepts from correct answer
        correct_words = set(correct_lower.split())
        user_words = set(user_lower.split())
        
        # Remove common words that don't carry meaning
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should'}
        
        meaningful_correct = correct_words - stop_words
        meaningful_user = user_words - stop_words
        
        # If user answer contains most key concepts, consider it correct
        if meaningful_correct and len(meaningful_user.intersection(meaningful_correct)) >= len(meaningful_correct) * 0.7:
            return True
            
        # Check for partial matches (if user answer is contained in or contains correct answer)
        if len(user_lower) > 5 and (user_lower in correct_lower or correct_lower in user_lower):
            return True
            
        return False

def chatbot_session(flashcards):
    """Interactive chatbot session using flashcards."""
    print("\n🤖 Flashcard Study Bot Started!")
    print("I'll ask you questions based on your flashcards.")
    print("Type 'quit' to exit, 'hint' for a hint, or 'skip' to skip.\n")
    
    # Ask user how many questions they want
    total_available = len(flashcards)
    print(f"Available flashcards: {total_available}")
    
    while True:
        try:
            desired_questions = input(f"How many questions do you want to answer? (1-{total_available}): ").strip()
            desired_questions = int(desired_questions)
            if 1 <= desired_questions <= total_available:
                break
            else:
                print(f"Please enter a number between 1 and {total_available}")
        except ValueError:
            print("Please enter a valid number")
    
    print(f"\nStarting study session with {desired_questions} questions...\n")
    
    score = 0
    questions_asked = 0
    
    # Shuffle flashcards for random order
    random.shuffle(flashcards)
    
    # Only take the number of questions requested
    selected_flashcards = flashcards[:desired_questions]
    
    for card in selected_flashcards:
        questions_asked += 1
        
        # Ask question based on card type
        if card['type'] == 'question_answer':
            question = card['question']
            correct_answer = card['answer']
        elif card['type'] == 'vocabulary':
            question = f"What is the definition of: {card['term']}?"
            correct_answer = card['definition']
        elif card['type'] == 'fact':
            question = card['prompt']
            correct_answer = card['content']
        else:
            continue
            
        print(f"📚 Question {questions_asked}/{desired_questions}:")
        print(f"Category: {card.get('category', 'General')} | Difficulty: {card.get('difficulty', 'Unknown')}")
        print(f"Q: {question}")
        
        # Get user response
        user_answer = input("\nYour answer: ").strip()
        
        if user_answer.lower() == 'quit':
            break
        elif user_answer.lower() == 'skip':
            print(f"⏭️  Skipped!")
            print(f"📖 Correct answer: {correct_answer}\n")
            continue
        elif user_answer.lower() == 'hint':
            # Generate hint using Claude
            hint = generate_hint(question, correct_answer)
            print(f"💡 Hint: {hint}")
            user_answer = input("Your answer (after hint): ").strip()
            
        # Check answer using Claude
        is_correct = check_answer(question, correct_answer, user_answer)
        
        if is_correct:
            print("✅ Correct!")
            score += 1
        else:
            print("❌ Incorrect.")
            
        # ALWAYS show the correct answer
        print(f"📖 Correct answer: {correct_answer}")
        print(f"Current score: {score}/{questions_asked}\n")
    
    # Final score with percentage
    if questions_asked > 0:
        percentage = (score / questions_asked) * 100
        print(f"\n🎉 Study session complete!")
        print(f"📊 Final Score: {score}/{questions_asked} ({percentage:.1f}%)")
        
        # Performance feedback
        if percentage >= 90:
            print("🌟 Excellent work! You've mastered this material!")
        elif percentage >= 80:
            print("👏 Great job! You have a strong understanding!")
        elif percentage >= 70:
            print("👍 Good work! Keep studying to improve!")
        elif percentage >= 60:
            print("📚 You're getting there! Review the material and try again!")
        else:
            print("💪 Keep studying! Practice makes perfect!")
    else:
        print("\n🎉 Study session ended. No questions were answered.")

def generate_hint(question, answer):
    """Generate a hint for the question using Claude."""
    prompt = f"""
    Generate a helpful hint for this study question without giving away the complete answer.
    
    Question: {question}
    Correct Answer: {answer}
    
    Provide a brief hint that guides the student toward the answer without revealing it completely.
    """
    
    try:
        response = client.messages.create(
            model="claude-3-7-sonnet-20250219",
            max_tokens=200,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        return response.content[0].text.strip()
    except:
        return "Think about the key concepts from your study material."

def main():
    """Main function to process PDF/URL and generate flashcards."""
    if len(sys.argv) < 2:
        print("Usage:")
        print("  Generate flashcards from document: python main.py <path_to_file>")
        print("  Generate flashcards from URL: python main.py <url>")
        print("  Study with chatbot: python main.py study <path_to_flashcard_json>")
        print("\nSupported document formats: PDF, DOC, DOCX")
        print("\nExamples:")
        print("  python main.py document.pdf")
        print("  python main.py document.docx")
        print("  python main.py https://en.wikipedia.org/wiki/Machine_Learning")
        print("  python main.py study document_flashcards.json")
        sys.exit(1)
    
    if sys.argv[1] == "study":
        # Chatbot mode
        if len(sys.argv) != 3:
            print("Usage for study mode: python main.py study <path_to_flashcard_json>")
            sys.exit(1)
            
        json_path = sys.argv[2]
        
        if not Path(json_path).exists():
            print(f"Error: Flashcard file '{json_path}' not found")
            sys.exit(1)
            
        print(f"Loading flashcards from: {json_path}")
        flashcards = load_flashcards(json_path)
        
        if not flashcards:
            print("Failed to load flashcards")
            sys.exit(1)
            
        print(f"Loaded {len(flashcards)} flashcards")
        chatbot_session(flashcards)
        
    else:
        # Content processing mode (PDF or URL)
        input_source = sys.argv[1]
        
        # Determine if input is URL or file path
        if input_source.startswith(('http://', 'https://')):
            # URL processing mode
            print(f"Processing URL: {input_source}")
            
            print("Extracting content from URL...")
            text = extract_text_from_url(input_source)
            
            if not text:
                print("Failed to extract content from URL")
                sys.exit(1)
                
            print(f"Extracted {len(text)} characters from URL")
            
            # Create filename from URL
            domain = urlparse(input_source).netloc.replace('.', '_')
            output_path = f"{domain}_flashcards.json"
            
        else:
            # Document processing mode (PDF, DOC, DOCX)
            file_path = input_source
            
            if not Path(file_path).exists():
                print(f"Error: File '{file_path}' not found")
                sys.exit(1)
            
            print(f"Processing document: {file_path}")
            
            print("Extracting text from document...")
            text = extract_text_from_document(file_path)
            
            if not text:
                print("Failed to extract text from document")
                sys.exit(1)
                
            print(f"Extracted {len(text)} characters from document")
            output_path = Path(file_path).stem + "_flashcards.json"
        
        # Generate flashcards using Claude (same for both sources)
        print("Generating flashcards with Claude AI...")
        flashcards = generate_flashcards(text)
        
        if not flashcards:
            print("Failed to generate flashcards")
            sys.exit(1)
        
        print(f"Generated {len(flashcards.get('flashcards', []))} flashcards")
        
        # Save flashcards to file
        save_flashcards(flashcards, output_path)
        
        # Also print to console
        print("\nGenerated Flashcards:")
        print(json.dumps(flashcards, indent=2, ensure_ascii=False))
        
        print(f"\n💡 To study with these flashcards, run:")
        print(f"python main.py study {output_path}")

def extract_text_from_docx(docx_path):
    """Extract text content from a DOCX file."""
    if not DOCX_AVAILABLE:
        raise Exception("python-docx not installed. Install with: pip install python-docx")
    
    try:
        doc = Document(docx_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return None

def extract_text_from_doc(doc_path):
    """Extract text content from a DOC file using textract."""
    try:
        import textract
        text = textract.process(doc_path).decode('utf-8')
        return text.strip()
    except ImportError:
        raise Exception("textract not installed. Install with: pip install textract")
    except Exception as e:
        print(f"Error reading DOC: {e}")
        return None

def extract_text_from_document(file_path):
    """Extract text from various document formats."""
    file_path = Path(file_path)
    extension = file_path.suffix.lower()
    
    if extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif extension == '.docx':
        return extract_text_from_docx(file_path)
    elif extension == '.doc':
        return extract_text_from_doc(file_path)
    else:
        raise Exception(f"Unsupported file format: {extension}. Supported formats: PDF, DOC, DOCX")

@app.post("/check-answer")
async def check_answer_endpoint(question: str, correct_answer: str, user_answer: str):
    is_correct = check_answer(question, correct_answer, user_answer)
    return {"correct": is_correct}

@app.post("/generate-from-url")
async def generate_flashcards_from_url(request: URLRequest):
    """API endpoint to generate flashcards from a URL."""
    try:
        import uuid
        
        print(f"Processing URL: {request.url}")
        text = extract_text_from_url(request.url)
        if not text:
            raise HTTPException(status_code=400, detail="Failed to extract content from URL")
        
        print(f"Extracted {len(text)} characters from URL")
        flashcards = generate_flashcards(text)
        if not flashcards:
            raise HTTPException(status_code=500, detail="Failed to generate flashcards")
        
        # Generate a session_id for compatibility with the frontend
        session_id = str(uuid.uuid4())
        
        return {
            "session_id": session_id,
            "flashcards": flashcards.get('flashcards', []),
            "message": f"Generated {len(flashcards.get('flashcards', []))} flashcards from URL"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing URL: {str(e)}")

if __name__ == "__main__":
    # CLI mode only - web server uses app.py
    main()